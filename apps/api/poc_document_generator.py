"""
Generate populated PoC / PoT Word documents from BigQuery-backed ProspectOverview.

Templates under templates/poc_pot/ use bracket placeholders like [Company Name] and
fixed boilerplate lines (trial dates, success examples). We discover those strings in
the document (body, tables, headers), extract structured facts via LLM from Gong + SFDC,
then fill deterministically (no second LLM pass for placeholders).

Key improvements (2026-05):
- Single structured AccountFacts extraction from Gong transcripts + SFDC
- Deterministic placeholder fills from extracted facts
- Decision Date from SFDC opp close_date or Gong timeline mentions
- Check-in Cadence defaults to "Weekly"
- Customer roster: names only, no "(N calls)" evidence text
- Arize roster: AE (opp owner) + SA always at top from SFDC
- Success criteria checkboxes: LLM picks 3-5 that resonate with Gong themes
- Confident, customer-facing tone (no hedging unless truly unknown)
"""

from __future__ import annotations

import io
import json
import os
import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Tuple

from openai_compat_completion import completion as llm_completion
from docx import Document
from docx.text.paragraph import Paragraph
from pydantic import BaseModel, Field, ValidationError

from models import ProspectOverview


class AppendixGenerationError(Exception):
    """Raised when the LLM output cannot be parsed into the expected JSON shape."""


BASE_DIR = Path(__file__).resolve().parent


def _find_template_dir() -> Path:
    """Find template directory - check api/ folder first (Vercel bundle), then root."""
    candidates = [
        BASE_DIR / "api" / "templates" / "poc_pot",
        BASE_DIR / "templates" / "poc_pot",
    ]
    for candidate in candidates:
        if candidate.is_dir() and any(candidate.glob("*.docx")):
            return candidate
    return candidates[-1]


TEMPLATE_DIR = _find_template_dir()

TEMPLATES: dict[str, Path] = {
    "poc_saas": TEMPLATE_DIR / "poc_saas.docx",
    "poc_vpc": TEMPLATE_DIR / "poc_vpc.docx",
    "pot": TEMPLATE_DIR / "pot.docx",
}

BRACKET_PLACEHOLDER_RE = re.compile(r"\[[^\]]{1,220}\]")

EXTRA_LITERAL_CANDIDATES_BY_TEMPLATE: dict[str, Tuple[str, ...]] = {
    "poc_saas": (
        "___________________________________________",
        "Start: _____ / End: _____ - [X] Weeks  ",
        "e.g., Mean time to root cause reduced from ___ days to < 4 hrs",
        "e.g., ___ prompt variants tested per sprint (target: 2-3x current)",
        "e.g., Drift alert configured for top 10 features; PSI threshold = ___",
    ),
    "poc_vpc": (
        "___________________________________________",
        "Start: _____ / End: _____ - [X] Weeks  ",
        "e.g., Mean time to root cause reduced from ___ days to < 4 hrs",
        "e.g., ___ prompt variants tested per sprint (target: 2-3x current)",
        "e.g., Drift alert configured for top 10 features; PSI threshold = ___",
    ),
    "pot": (
        "Start: _____ / End: _____ - [X] Weeks  ",
    ),
}


# =============================================================================
# AccountFacts: Structured extraction from Gong + SFDC
# =============================================================================


class AccountFacts(BaseModel):
    """Structured facts extracted from Gong transcripts + SFDC for PoC/PoT docs.

    All fields are optional; missing = unknown (dash in output).
    """

    # Tech stack (from Gong discussions)
    framework: Optional[str] = Field(
        None, description="Orchestration framework: LangChain, LangGraph, LlamaIndex, Custom, etc."
    )
    current_state_observability: Optional[str] = Field(
        None, description="How they do observability/evals today: DataDog, Splunk, Homegrown, None"
    )
    llm_provider: Optional[str] = Field(
        None, description="LLM provider(s): OpenAI, Anthropic, AWS Bedrock, GCP Vertex, Azure OAI"
    )
    vector_db: Optional[str] = Field(
        None, description="Vector DB if RAG: Pinecone, Weaviate, pgvector, Internal, N/A"
    )
    coding_agent: Optional[str] = Field(
        None, description="Predominant coding agent: Cursor, Claude Code, Copilot, Windsurf, etc."
    )
    app_type: Optional[str] = Field(
        None, description="Type of application: LLM Calls, RAG, Agentic (Tool Calls), Multi-Agent"
    )

    # Use case - the customer's SPECIFIC AI application (NOT Arize's product, NOT generic)
    application_name: Optional[str] = Field(
        None,
        description="The SPECIFIC AI agent or LLM app the customer is building - e.g., 'Crypto Price Prediction Agent', "
        "'Customer Support Chatbot', 'Document Summarizer'. Must be THEIR product. Return null if not discussed.",
    )
    use_case_summary: Optional[str] = Field(
        None,
        description="What does the customer's specific AI app do? Must describe THEIR product. Null if not discussed.",
    )
    expected_trace_volume: Optional[str] = Field(
        None, description="Expected spans/day if mentioned (e.g., '~50K spans/day')"
    )
    production_or_staging: Optional[str] = Field(
        None, description="Production, Staging, or Sample App"
    )

    # Timeline
    decision_date: Optional[str] = Field(
        None, description="Target decision date from Gong or SFDC close_date (YYYY-MM-DD)"
    )
    timeline_context: Optional[str] = Field(
        None, description="Any timeline/urgency context from Gong calls"
    )

    # Success criteria checkboxes (row indices that resonate with Gong themes)
    success_criteria_indices: List[int] = Field(
        default_factory=list,
        description="Indices (0-based) of success criteria rows to check (3-5 recommended)",
    )

    # Custom success criteria targets (for e.g., ___ lines)
    success_target_mttr: Optional[str] = Field(
        None, description="Custom target for MTTR line, or N/A if not applicable"
    )
    success_target_prompts: Optional[str] = Field(
        None, description="Custom target for prompt variants line, or N/A if not applicable"
    )
    success_target_drift: Optional[str] = Field(
        None, description="Custom target for drift monitoring line, or N/A if not applicable"
    )

    # One-line opportunity summary (for the long underline placeholder)
    opportunity_summary: Optional[str] = Field(
        None, description="One-line summary of how they measure success today"
    )


class IntegratedFillResult(BaseModel):
    """LLM output: map each template token to its replacement string."""

    replacements: Dict[str, str] = Field(default_factory=dict)


def _deep_truncate(obj: Any, max_str: int = 2200, max_list: int = 60) -> Any:
    if isinstance(obj, dict):
        return {k: _deep_truncate(v, max_str, max_list) for k, v in obj.items()}
    if isinstance(obj, list):
        return [_deep_truncate(i, max_str, max_list) for i in obj[:max_list]]
    if isinstance(obj, str) and len(obj) > max_str:
        return obj[:max_str] + "…"
    return obj


def _strip_json_fence(text: str) -> str:
    t = text.strip()
    if t.startswith("```"):
        parts = t.split("```")
        if len(parts) >= 2:
            inner = parts[1]
            if inner.lstrip().startswith("json"):
                inner = inner.lstrip()[4:]
            return inner.strip()
    return t


def _extract_json_object(raw: str) -> dict:
    """Parse a JSON object from an LLM message that may include prose or markdown fences."""
    s = raw.strip().lstrip("\ufeff")
    s = _strip_json_fence(s)
    try:
        data = json.loads(s)
        if isinstance(data, dict):
            return data
    except json.JSONDecodeError:
        pass
    start = s.find("{")
    if start < 0:
        raise ValueError("No JSON object found in model response")
    decoder = json.JSONDecoder()
    try:
        obj, _end = decoder.raw_decode(s[start:])
    except json.JSONDecodeError as e:
        raise ValueError(f"Invalid JSON in model response: {e}") from e
    if not isinstance(obj, dict):
        raise ValueError("Top-level JSON must be an object")
    return obj


def _overview_blob(overview: ProspectOverview) -> str:
    try:
        return json.dumps(overview.model_dump(mode="json", exclude_none=True), default=str).lower()
    except Exception:
        return ""


def _build_gong_context_for_extraction(overview: ProspectOverview) -> str:
    """Build rich Gong context for the AccountFacts extraction prompt."""
    sections: list[str] = []
    gong = overview.gong_summary
    if not gong or not gong.recent_calls:
        return "No Gong calls available."

    for call in gong.recent_calls[:12]:
        title = (call.call_title or "Untitled").strip()
        date = (call.call_date or "")[:10] or "?"
        lines = [f"### {date} — {title}"]

        if call.spotlight_brief:
            lines.append(f"**Brief:** {call.spotlight_brief[:500]}")
        if call.spotlight_key_points:
            lines.append("**Key points:** " + " | ".join(str(k)[:200] for k in call.spotlight_key_points[:8]))
        if call.spotlight_next_steps:
            lines.append(f"**Next steps:** {call.spotlight_next_steps[:300]}")
        if call.spotlight_outcome:
            lines.append(f"**Outcome:** {call.spotlight_outcome[:200]}")

        if call.transcript_snippet:
            snippet = call.transcript_snippet if isinstance(call.transcript_snippet, str) else str(call.transcript_snippet)
            lines.append(f"**Transcript excerpt:** {snippet[:800]}")

        sections.append("\n".join(lines))

    if gong.key_themes:
        sections.append("**Extracted themes:** " + ", ".join(gong.key_themes[:15]))

    return "\n\n".join(sections)


def _build_sfdc_context_for_extraction(overview: ProspectOverview) -> str:
    """Build SFDC context for the AccountFacts extraction prompt."""
    bits: list[str] = []
    sf = overview.salesforce
    if sf:
        if sf.name:
            bits.append(f"Account: {sf.name}")
        if sf.industry:
            bits.append(f"Industry: {sf.industry}")
        if sf.is_using_llms:
            bits.append(f"Using LLMs: {sf.is_using_llms}")
        if sf.deployment_types:
            bits.append(f"Deployment: {sf.deployment_types}")
        if sf.description:
            bits.append(f"Description: {sf.description[:400]}")
        if sf.customer_notes:
            bits.append(f"Notes: {sf.customer_notes[:400]}")

    opp = overview.latest_opportunity
    if opp:
        if opp.name:
            bits.append(f"Opportunity: {opp.name}")
        if opp.stage_name:
            bits.append(f"Stage: {opp.stage_name}")
        if opp.close_date:
            bits.append(f"Close date: {opp.close_date}")
        if opp.description:
            bits.append(f"Opp description: {opp.description[:500]}")
        if opp.next_step:
            bits.append(f"Next step: {opp.next_step[:300]}")

    return "\n".join(bits) if bits else "No SFDC data available."


def _extract_account_facts(
    overview: ProspectOverview,
    document_template: str,
    manual_notes: str | None,
    llm_model: str | None = None,
) -> AccountFacts:
    """Single LLM call to extract structured AccountFacts from Gong + SFDC."""
    model = llm_model or os.environ.get("POC_DOC_LLM_MODEL", "claude-haiku-4-5")

    gong_context = _build_gong_context_for_extraction(overview)
    sfdc_context = _build_sfdc_context_for_extraction(overview)

    notes_block = ""
    if (manual_notes or "").strip():
        notes_block = f"\n## SA Notes (high priority)\n{manual_notes.strip()}\n"

    # Derive SFDC close_date as fallback for decision_date
    sfdc_close_date = None
    if overview.latest_opportunity and overview.latest_opportunity.close_date:
        sfdc_close_date = overview.latest_opportunity.close_date[:10]

    prompt = f"""You are extracting structured facts from Gong call transcripts and Salesforce data to populate an Arize AX PoC/PoT document.

## Document type: {document_template.upper()}

## Gong Calls
{gong_context}

## Salesforce Data
{sfdc_context}
{notes_block}

## Your task
Return a JSON object with the following fields. Use ONLY facts grounded in the data above.
If something is not mentioned, use null (not "TBD" or "unknown").

{{
  "framework": "LangChain | LangGraph | LlamaIndex | Custom | null",
  "current_state_observability": "DataDog | Splunk | Homegrown | None | null (how they do obs/evals today)",
  "llm_provider": "OpenAI | Anthropic | AWS Bedrock | GCP Vertex | Azure OAI | comma-separated if multiple | null",
  "vector_db": "Pinecone | Weaviate | pgvector | Internal | N/A | null",
  "coding_agent": "Cursor | Claude Code | Copilot | Windsurf | null (if they mention using coding assistants)",
  "app_type": "LLM Calls | RAG | Agentic (Tool Calls) | Multi-Agent | null",
  "application_name": "The customer's SPECIFIC AI app name - e.g., 'Crypto Research Agent', 'Support Chatbot', 'Doc Summarizer'. WRONG answers: 'uses Arize for...', 'AI observability...', 'tracing agentic workflows'. If calls only discuss pricing/contracts/security and no specific app is named, return null. | null",
  "use_case_summary": "What specific AI app is the customer building? WRONG: anything mentioning Arize/observability/tracing. Return null if not explicitly discussed. | null",
  "expected_trace_volume": "~50K spans/day | null (if mentioned)",
  "production_or_staging": "Production | Staging | Sample App | null",
  "decision_date": "YYYY-MM-DD | null (from timeline discussions or SFDC close_date: {sfdc_close_date or 'not set'})",
  "timeline_context": "any urgency or timeline context from calls | null",
  "success_criteria_indices": [list of 0-based row indices from success criteria table that match their needs, pick 3-5],
  "success_target_mttr": "custom target for MTTR if applicable, else 'N/A — not discussed'",
  "success_target_prompts": "custom target for prompt iteration if applicable, else 'N/A — not discussed'",
  "success_target_drift": "custom target for drift monitoring if applicable, else 'N/A — not discussed'",
  "opportunity_summary": "one-line summary of how they measure success today or what problem they're solving"
}}

## Success criteria rows (for success_criteria_indices):
Row 0: LLM / Agent Observability header
Row 1: Framework-agnostic tracing via OpenInference/OTel
Row 2: Session-level agent evaluation
Row 3: Online LLM-as-a-judge evals (hallucination, relevance, Q&A correctness)
Row 4: AI copilot autonomous root-cause analysis (Alyx)
Row 5: Prompt IDE with version control
Row 6: Export curated datasets from traces for experiments
Row 7: Production dashboards (cost, latency, token usage)
Row 8: Production-grade alerting (Slack, PagerDuty)
Row 9: Headless debugging via CLI and Skills
Row 10: Traditional ML Observability header
Row 11: Unified LLM + traditional ML observability
Row 12: Feature-level drift monitoring
Row 13: Data quality monitoring
Row 14: Custom metrics and monitoring rules
Row 15: Platform/Security header
Row 16: Scalability (large data volumes)
Row 17: Flexible deployment (SaaS, VPC, self-hosted, CMEK)
Row 18: SOC 2 + HIPAA certified, SSO/SAML, audit logs

Pick 3-5 rows that ACTUALLY match their discussed use cases and pain points. Skip headers (0, 10, 15).

Return ONLY the JSON object, no explanation.
"""

    response = llm_completion(
        model=model,
        messages=[{"role": "user", "content": prompt}],
        max_tokens=2048,
        temperature=0.1,
    )
    raw = (response.choices[0].message.content or "").strip()

    if os.environ.get("POC_DOC_DEBUG", "").strip().lower() in ("1", "true", "yes"):
        print(f"[poc_document_generator] AccountFacts raw: {raw[:1000]!r}")

    try:
        data = _extract_json_object(raw)
        facts = AccountFacts.model_validate(data)
        
        # Post-process: clear application_name/use_case_summary if they describe Arize's product
        bad_patterns = [
            "arize", "observability", "tracing", "monitoring", "spans", 
            "annotation", "agentic workflow", "ai observability", "llm observability",
            "agentic ai", "requires tracing", "building agentic", "agentic systems",
            "evaluation", "evals", "llm ops", "mlops"
        ]
        
        # Filter application_name
        if facts.application_name:
            app_lower = facts.application_name.lower()
            if any(bp in app_lower for bp in bad_patterns):
                facts.application_name = None
        
        # Filter use_case_summary too
        if facts.use_case_summary:
            summary_lower = facts.use_case_summary.lower()
            if any(bp in summary_lower for bp in bad_patterns):
                facts.use_case_summary = None
        
        return facts
    except (ValueError, json.JSONDecodeError, ValidationError) as e:
        if os.environ.get("POC_DOC_DEBUG"):
            print(f"[poc_document_generator] AccountFacts parse error: {e}")
        return AccountFacts()


def _company_display_name(overview: ProspectOverview) -> str:
    if overview.salesforce and (overview.salesforce.name or "").strip():
        return overview.salesforce.name.strip()
    if (overview.lookup_value or "").strip():
        return overview.lookup_value.strip()
    return "Customer"


def _get_ae_name(overview: ProspectOverview) -> Optional[str]:
    """Get AE name from SFDC opportunity owner or account owner."""
    if overview.latest_opportunity and overview.latest_opportunity.owner_name:
        return overview.latest_opportunity.owner_name.strip()
    if overview.salesforce and overview.salesforce.owner_name:
        return overview.salesforce.owner_name.strip()
    return None


def _get_sa_name(overview: ProspectOverview) -> Optional[str]:
    """Get SA name from SFDC account assigned_sa."""
    if overview.salesforce and overview.salesforce.assigned_sa:
        return overview.salesforce.assigned_sa.strip()
    return None


def _build_placeholder_map_from_facts(
    facts: AccountFacts,
    overview: ProspectOverview,
    document_template: str,
) -> Dict[str, str]:
    """Build deterministic placeholder replacements from extracted AccountFacts."""
    company = _company_display_name(overview)
    m: Dict[str, str] = {}

    # Company name
    m["[Company Name]"] = company

    # Duration weeks
    weeks = {"pot": "1", "poc_saas": "2", "poc_vpc": "4"}.get(document_template, "2")
    m["[X]"] = weeks

    # Start/End dates
    m["Start: _____ / End: _____ - [X] Weeks  "] = f"Start: TBD / End: TBD — {weeks} weeks  "

    # Framework
    if facts.framework:
        m["[LangChain / LangGraph / LlamaIndex / Custom / etc.]"] = facts.framework
        m["[LangChain / LlamaIndex / LangGraph / Custom / None]"] = facts.framework
    else:
        m["[LangChain / LangGraph / LlamaIndex / Custom / etc.]"] = "—"
        m["[LangChain / LlamaIndex / LangGraph / Custom / None]"] = "—"

    # Current state (observability)
    if facts.current_state_observability:
        m["[DataDog / Splunk / Homegrown / None]"] = facts.current_state_observability
    else:
        m["[DataDog / Splunk / Homegrown / None]"] = "—"

    # Environment (derive from template type)
    env_map = {"poc_saas": "SaaS", "poc_vpc": "VPC", "pot": "SaaS"}
    m["[SaaS / VPC / On-Prem]"] = env_map.get(document_template, "SaaS")

    # Coding agent
    if facts.coding_agent:
        m["[e.g., Claude Code, Cursor, etc.]"] = facts.coding_agent
    else:
        m["[e.g., Claude Code, Cursor, etc.]"] = "—"

    # Application name/description
    if facts.application_name:
        m["[e.g., Customer Support Bot, Surveillance Agent]"] = facts.application_name
    elif facts.use_case_summary:
        m["[e.g., Customer Support Bot, Surveillance Agent]"] = facts.use_case_summary[:80]
    else:
        m["[e.g., Customer Support Bot, Surveillance Agent]"] = "—"

    # Expected trace volume
    if facts.expected_trace_volume:
        m["[e.g., ~50K spans/day — estimate OK to start]"] = facts.expected_trace_volume
    else:
        m["[e.g., ~50K spans/day — estimate OK to start]"] = "—"

    # LLM Provider
    if facts.llm_provider:
        m["[OpenAI / Anthropic / AWS Bedrock / GCP Vertex / Azure OAI]"] = facts.llm_provider
    else:
        m["[OpenAI / Anthropic / AWS Bedrock / GCP Vertex / Azure OAI]"] = "—"

    # Vector DB
    if facts.vector_db:
        m["[Pinecone / Weaviate / pgvector / Internal / N/A]"] = facts.vector_db
    else:
        m["[Pinecone / Weaviate / pgvector / Internal / N/A]"] = "N/A"

    # App type (PoT only)
    if facts.app_type:
        m["[LLM Calls / RAG / Agentic (Tool Calls) / Multi-Agent]"] = facts.app_type
    else:
        m["[LLM Calls / RAG / Agentic (Tool Calls) / Multi-Agent]"] = "—"

    # Opportunity summary (long underline)
    if facts.opportunity_summary:
        m["___________________________________________"] = facts.opportunity_summary
    else:
        m["___________________________________________"] = f"{company} — evaluating Arize AX for AI observability"

    # Success criteria custom targets
    if facts.success_target_mttr:
        m["e.g., Mean time to root cause reduced from ___ days to < 4 hrs"] = facts.success_target_mttr
    else:
        m["e.g., Mean time to root cause reduced from ___ days to < 4 hrs"] = "N/A — not discussed"

    if facts.success_target_prompts:
        m["e.g., ___ prompt variants tested per sprint (target: 2-3x current)"] = facts.success_target_prompts
    else:
        m["e.g., ___ prompt variants tested per sprint (target: 2-3x current)"] = "N/A — not discussed"

    if facts.success_target_drift:
        m["e.g., Drift alert configured for top 10 features; PSI threshold = ___"] = facts.success_target_drift
    else:
        m["e.g., Drift alert configured for top 10 features; PSI threshold = ___"] = "N/A — not discussed"

    return m


def _iter_paragraphs_in_cell(cell: Any) -> Iterable[Paragraph]:
    for p in cell.paragraphs:
        yield p
    for t in cell.tables:
        for row in t.rows:
            for c in row.cells:
                yield from _iter_paragraphs_in_cell(c)


def _iter_body_paragraphs(doc: Document) -> Iterable[Paragraph]:
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                yield from _iter_paragraphs_in_cell(cell)


def _iter_part_paragraphs(part: Any) -> Iterable[Paragraph]:
    if part is None:
        return
    for p in part.paragraphs:
        yield p
    for t in part.tables:
        for row in t.rows:
            for cell in row.cells:
                yield from _iter_paragraphs_in_cell(cell)


def iter_all_paragraphs(doc: Document) -> Iterable[Paragraph]:
    """Body, tables, and header/footer paragraphs."""
    yield from _iter_body_paragraphs(doc)
    for section in doc.sections:
        yield from _iter_part_paragraphs(section.header)
        yield from _iter_part_paragraphs(section.footer)
        if section.different_first_page_header_footer:
            yield from _iter_part_paragraphs(section.first_page_header)
            yield from _iter_part_paragraphs(section.first_page_footer)


def discover_replacement_keys(doc: Document, document_template: str) -> List[str]:
    """Collect bracket tokens plus any configured literal line that exists in this file."""
    found: set[str] = set()
    paras = list(iter_all_paragraphs(doc))
    for p in paras:
        found.update(BRACKET_PLACEHOLDER_RE.findall(p.text))
    for literal in EXTRA_LITERAL_CANDIDATES_BY_TEMPLATE.get(document_template, ()):
        for p in paras:
            if literal in p.text:
                found.add(literal)
                break
    # Longest keys first so shorter tokens (e.g. [X]) do not break longer literals prematurely.
    return sorted(found, key=len, reverse=True)


def _set_paragraph_text_collapsed(paragraph: Paragraph, new_text: str) -> None:
    runs = paragraph.runs
    if runs:
        runs[0].text = new_text
        for r in runs[1:]:
            r.text = ""
    else:
        paragraph.add_run(new_text)


def apply_replacements(doc: Document, mapping: List[Tuple[str, str]]) -> None:
    """Apply (old, new) pairs to every paragraph; longest keys should appear first in the list."""
    for paragraph in iter_all_paragraphs(doc):
        text = paragraph.text
        orig = text
        for old, new in mapping:
            if old in text:
                text = text.replace(old, new)
        if text != orig:
            _set_paragraph_text_collapsed(paragraph, text)


def _set_cell_text(cell: Any, text: str) -> None:
    """Write plain text into the first paragraph of a table cell."""
    if cell.paragraphs:
        _set_paragraph_text_collapsed(cell.paragraphs[0], text)
    else:
        cell.add_paragraph(text)


def _cell_is_effectively_empty(cell: Any) -> bool:
    return not any(p.text.strip() for p in cell.paragraphs)


def _row_data_cells_empty(row: Any) -> bool:
    return all(_cell_is_effectively_empty(c) for c in row.cells)


def _table_header_cells(row: Any) -> List[str]:
    return [c.text.strip() for c in row.cells]


def _rank_participants_across_calls(
    overview: ProspectOverview,
    affiliation_filter: str,
) -> List[Tuple[str, str, int]]:
    """Rank participants by frequency + recency across ALL Gong calls.

    affiliation_filter: "customer" to get external/prospect participants,
                        "arize" to get internal/company participants.

    Returns list of (name, email, call_count) sorted by call_count desc, then
    by most-recent appearance.
    """
    gong = overview.gong_summary
    if not gong:
        return []

    is_customer = affiliation_filter == "customer"
    counts: Dict[str, Dict[str, Any]] = {}

    for idx, call in enumerate(gong.recent_calls):
        for p in call.participants:
            name = (p.name or "").strip()
            if not name:
                continue
            aff = (p.affiliation or "").strip().lower()

            if is_customer:
                if aff in ("company", "internal"):
                    continue
                if aff not in ("external", "customer", "prospect", "non_company", ""):
                    continue
            else:
                if aff != "company":
                    continue

            key = name.lower()
            if key not in counts:
                counts[key] = {
                    "name": name,
                    "email": (p.email or "").strip(),
                    "call_count": 0,
                    "first_seen_idx": idx,
                }
            counts[key]["call_count"] += 1
            if (p.email or "").strip() and not counts[key]["email"]:
                counts[key]["email"] = (p.email or "").strip()

    ranked = sorted(
        counts.values(),
        key=lambda x: (-x["call_count"], x["first_seen_idx"]),
    )
    return [(r["name"], r["email"], r["call_count"]) for r in ranked]


def _ranked_customer_participants(overview: ProspectOverview) -> List[Tuple[str, str, int]]:
    return _rank_participants_across_calls(overview, "customer")


def _ranked_arize_participants(overview: ProspectOverview) -> List[Tuple[str, str, int]]:
    return _rank_participants_across_calls(overview, "arize")


def _arize_roster_lines(overview: ProspectOverview) -> List[Tuple[str, str, str, str]]:
    """Build ordered list of (name, title, role, email) for Arize team members.

    Sources in priority order (AE + SA always at top):
    1. AE from SFDC opportunity owner (or account owner)
    2. SA from SFDC assigned_sa
    3. Other SFDC team assignments (CSE, AI SE, CSM)
    4. Arize participants from Gong calls
    Deduplicates by lowered name.
    """
    seen: set[str] = set()
    lines: List[Tuple[str, str, str, str]] = []

    # AE first (from opp owner or account owner)
    ae_name = _get_ae_name(overview)
    if ae_name and ae_name.lower() not in seen:
        seen.add(ae_name.lower())
        lines.append((ae_name, "Account Executive", "Arize", ""))

    # SA second
    sa_name = _get_sa_name(overview)
    if sa_name and sa_name.lower() not in seen:
        seen.add(sa_name.lower())
        lines.append((sa_name, "Solution Architect", "Arize", ""))

    # Other SFDC assignments
    if overview.salesforce:
        sf = overview.salesforce
        other_assignments = [
            ((getattr(sf, "assigned_cse", None) or "").strip(), "Customer Success Engineer"),
            ((getattr(sf, "assigned_ai_se", None) or "").strip(), "AI Sales Engineer"),
            ((getattr(sf, "assigned_csm", None) or "").strip(), "Customer Success Manager"),
        ]
        for name, title in other_assignments:
            if name and name.lower() not in seen:
                seen.add(name.lower())
                lines.append((name, title, "Arize", ""))

    # Gong participants (without call count in title)
    for name, email, _count in _ranked_arize_participants(overview):
        if name.lower() not in seen:
            seen.add(name.lower())
            lines.append((name, "Arize Team", "Arize", email))

    return lines


def _fill_empty_data_rows(
    table: Any,
    header_check: List[str],
    rows_data: List[List[str]],
) -> None:
    """Fill empty data rows in a Word table, skipping the header row."""
    if _table_header_cells(table.rows[0]) != header_check:
        return
    data_row_idx = 1
    for values in rows_data:
        if data_row_idx >= len(table.rows):
            break
        if _row_data_cells_empty(table.rows[data_row_idx]):
            for ci, val in enumerate(values):
                if ci < len(table.rows[data_row_idx].cells):
                    _set_cell_text(table.rows[data_row_idx].cells[ci], val[:240])
            data_row_idx += 1
        else:
            data_row_idx += 1


def fill_structured_roster_cells(
    doc: Document,
    overview: ProspectOverview,
    document_template: str,
    facts: Optional[AccountFacts] = None,
) -> None:
    """Fill roster, kickoff tables, and success criteria checkboxes.

    Customer roster: names only (no call count evidence).
    Arize roster: AE + SA always at top from SFDC.
    Also fills: Decision Date, Check-in/Workshop Cadence (Weekly), escalation contacts,
    and success criteria checkboxes based on AccountFacts.
    """
    customer_ranked = _ranked_customer_participants(overview)
    arize_lines = _arize_roster_lines(overview)
    ae_name = _get_ae_name(overview)
    top_customer = customer_ranked[0][0] if customer_ranked else None

    # --- Decision Date: fill in table 0 for all templates ---
    if len(doc.tables) > 0:
        t0 = doc.tables[0]
        for ri in range(len(t0.rows)):
            label = t0.rows[ri].cells[0].text.strip()
            if "Decision Date" in label and _cell_is_effectively_empty(t0.rows[ri].cells[1]):
                decision_date = None
                if facts and facts.decision_date:
                    decision_date = facts.decision_date
                elif overview.latest_opportunity and overview.latest_opportunity.close_date:
                    decision_date = overview.latest_opportunity.close_date[:10]
                if decision_date:
                    _set_cell_text(t0.rows[ri].cells[1], decision_date)

    # --- PoT-specific: table 1 has Item | Details (Slack, Workshop Cadence) ---
    if document_template == "pot" and len(doc.tables) > 1:
        t1 = doc.tables[1]
        if _table_header_cells(t1.rows[0]) == ["Item", "Details"]:
            for ri in range(1, len(t1.rows)):
                label = t1.rows[ri].cells[0].text.strip()
                # Workshop Cadence defaults to Weekly
                if "Workshop Cadence" in label and _cell_is_effectively_empty(t1.rows[ri].cells[1]):
                    _set_cell_text(t1.rows[ri].cells[1], "Weekly")
        return  # PoT has no roster tables

    # --- PoC templates only below ---
    if document_template not in ("poc_saas", "poc_vpc") or len(doc.tables) < 4:
        return

    # --- Customer roster (table 1): Name | Role / Team | % Time Committed | Responsibilities ---
    if customer_ranked:
        customer_rows = []
        for name, email, _count in customer_ranked[:6]:
            customer_rows.append([name, "Customer", "—", email or "—"])
        _fill_empty_data_rows(
            doc.tables[1],
            ["Name", "Role / Team", "% Time Committed", "Responsibilities"],
            customer_rows,
        )

    # --- Arize roster (table 2): Name | Title | Email | Role ---
    if arize_lines:
        arize_rows = [[nm, title, email, role] for nm, title, role, email in arize_lines[:6]]
        _fill_empty_data_rows(
            doc.tables[2],
            ["Name", "Title", "Email", "Role"],
            arize_rows,
        )

    # --- Kickoff/Communication table (table 3): Item | Details ---
    t3 = doc.tables[3]
    if _table_header_cells(t3.rows[0]) == ["Item", "Details"]:
        for ri in range(1, len(t3.rows)):
            label = t3.rows[ri].cells[0].text.strip()

            # Check-in Cadence defaults to Weekly
            if "Check-in Cadence" in label and _cell_is_effectively_empty(t3.rows[ri].cells[1]):
                _set_cell_text(t3.rows[ri].cells[1], "Weekly")

            # Primary Escalation Contact (Arize) = AE
            if "Primary Escalation Contact (Arize)" in label and _cell_is_effectively_empty(t3.rows[ri].cells[1]):
                if ae_name:
                    _set_cell_text(t3.rows[ri].cells[1], ae_name[:240])

            # Primary Escalation Contact (Customer) = top customer participant
            if "Primary Escalation Contact (Customer)" in label and _cell_is_effectively_empty(t3.rows[ri].cells[1]):
                if top_customer:
                    _set_cell_text(t3.rows[ri].cells[1], top_customer[:240])

    # --- Success criteria checkboxes (table 7 for poc_saas, table 8 for poc_vpc) ---
    success_table_idx = 7 if document_template == "poc_saas" else 8
    if len(doc.tables) > success_table_idx and facts and facts.success_criteria_indices:
        success_table = doc.tables[success_table_idx]
        for row_idx in facts.success_criteria_indices:
            if 0 < row_idx < len(success_table.rows):
                checkbox_cell = success_table.rows[row_idx].cells[0]
                if _cell_is_effectively_empty(checkbox_cell):
                    _set_cell_text(checkbox_cell, "☒")


def generate_integrated_fills(
    *,
    keys: List[str],
    overview: ProspectOverview,
    document_template: str,
    manual_notes: str | None,
    facts: AccountFacts,
    llm_model: str | None = None,
) -> Dict[str, str]:
    """Generate placeholder fills using pre-extracted AccountFacts.

    Most bracket placeholders are filled deterministically from facts.
    Only falls back to LLM for keys not covered by facts.
    """
    # Build deterministic map from facts
    facts_map = _build_placeholder_map_from_facts(facts, overview, document_template)

    # Merge facts_map with discovered keys
    merged: Dict[str, str] = {}
    for k in keys:
        if k in facts_map and facts_map[k] and facts_map[k] != "—":
            merged[k] = facts_map[k]
        elif k == "[Company Name]":
            merged[k] = _company_display_name(overview)
        elif k == "[X]":
            merged[k] = {"pot": "1", "poc_saas": "2", "poc_vpc": "4"}.get(document_template, "2")
        elif k.startswith("[") and k.endswith("]"):
            merged[k] = facts_map.get(k, "—")
        elif k.startswith("e.g.,"):
            merged[k] = facts_map.get(k, "N/A — not discussed")
        elif k == "___________________________________________":
            merged[k] = facts_map.get(k, f"{_company_display_name(overview)} — evaluating Arize AX")
        elif k.startswith("Start: _____"):
            weeks = {"pot": "1", "poc_saas": "2", "poc_vpc": "4"}.get(document_template, "2")
            merged[k] = f"Start: TBD / End: TBD — {weeks} weeks  "
        else:
            merged[k] = facts_map.get(k, "—")

    return merged


def _safe_filename_part(name: str) -> str:
    s = "".join(c for c in name if c.isalnum() or c in (" ", "-", "_")).strip()
    return s.replace(" ", "_")[:80] or "Account"


# Human-readable segment for download filenames (filesystem-safe, no slashes).
DOCUMENT_EXPORT_TYPE_LABELS: dict[str, str] = {
    "poc_saas": "PoC-SaaS",
    "poc_vpc": "PoC-VPC",
    "pot": "PoT",
}


def _export_type_filename_part(document_template: str) -> str:
    label = DOCUMENT_EXPORT_TYPE_LABELS.get(document_template, document_template)
    return _safe_filename_part(label)


def build_poc_document(
    overview: ProspectOverview,
    document_template: str,
    manual_notes: str | None = None,
    llm_model: str | None = None,
) -> tuple[bytes, str]:
    """Build a populated PoC/PoT Word document.

    1. Extract structured AccountFacts from Gong + SFDC via single LLM call
    2. Fill bracket placeholders deterministically from facts
    3. Fill roster tables (customer names only, AE+SA at top for Arize)
    4. Fill Decision Date, Check-in Cadence (Weekly), escalation contacts
    5. Check success criteria boxes based on Gong themes
    """
    path = TEMPLATES.get(document_template)
    if not path or not path.is_file():
        raise FileNotFoundError(f"Missing template for {document_template}: {path}")

    doc = Document(str(path))
    keys = discover_replacement_keys(doc, document_template)
    if not keys:
        raise ValueError("No replaceable placeholders found in template")

    # Step 1: Extract structured facts from Gong + SFDC
    facts = _extract_account_facts(
        overview=overview,
        document_template=document_template,
        manual_notes=manual_notes,
        llm_model=llm_model,
    )

    # Step 2: Build placeholder fills from facts
    fill_map = generate_integrated_fills(
        keys=keys,
        overview=overview,
        document_template=document_template,
        manual_notes=manual_notes,
        facts=facts,
        llm_model=llm_model,
    )
    ordered = sorted(fill_map.items(), key=lambda kv: len(kv[0]), reverse=True)
    apply_replacements(doc, ordered)

    # Step 3-5: Fill rosters, kickoff table, success criteria checkboxes
    fill_structured_roster_cells(doc, overview, document_template, facts)

    buf = io.BytesIO()
    doc.save(buf)
    date_part = datetime.now(timezone.utc).strftime("%Y%m%d")
    account_safe = _safe_filename_part(_company_display_name(overview))
    type_safe = _export_type_filename_part(document_template)
    fname = f"Arize_AX_{account_safe}_{type_safe}_{date_part}.docx"
    return buf.getvalue(), fname
