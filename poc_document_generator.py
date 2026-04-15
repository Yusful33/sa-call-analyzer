"""
Generate populated PoC / PoT Word documents from BigQuery-backed ProspectOverview.

Templates under templates/poc_pot/ use bracket placeholders like [Company Name] and
fixed boilerplate lines (trial dates, success examples). We discover those strings in
the document (body, tables, headers), ask an LLM for replacement text per key, then
substitute in-place (no appendix appended at the end).
"""

from __future__ import annotations

import io
import json
import os
import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, Iterable, List, Tuple

import litellm
from docx import Document
from docx.text.paragraph import Paragraph
from pydantic import BaseModel, Field, ValidationError

from models import ProspectOverview


class AppendixGenerationError(Exception):
    """Raised when the LLM output cannot be parsed into the expected JSON shape."""


BASE_DIR = Path(__file__).resolve().parent
TEMPLATE_DIR = BASE_DIR / "templates" / "poc_pot"

TEMPLATES: dict[str, Path] = {
    "poc_saas": TEMPLATE_DIR / "poc_saas.docx",
    "poc_vpc": TEMPLATE_DIR / "poc_vpc.docx",
    "pot": TEMPLATE_DIR / "pot.docx",
}

BRACKET_PLACEHOLDER_RE = re.compile(r"\[[^\]]{1,220}\]")

# Long literal lines present in PoC templates (not always in PoT). We only require keys
# that actually appear somewhere in the loaded document.
EXTRA_LITERAL_CANDIDATES_BY_TEMPLATE: dict[str, Tuple[str, ...]] = {
    "poc_saas": (
        "___________________________________________",
        "Start: _____ / End: _____ - [X] Weeks  ",
        "e.g., Mean time to root cause reduced from ___ days to < 4 hrs",
        "e.g., ___ prompt variants tested per sprint (target: 2-3x current)",
        "e.g., Drift alert configured for top 10 features; PSI threshold = ___",
    ),
    "poc_vpc": (
        "___________________________________________",
        "Start: _____ / End: _____ - [X] Weeks  ",
        "e.g., Mean time to root cause reduced from ___ days to < 4 hrs",
        "e.g., ___ prompt variants tested per sprint (target: 2-3x current)",
        "e.g., Drift alert configured for top 10 features; PSI threshold = ___",
    ),
    "pot": (
        "Start: _____ / End: _____ - [X] Weeks  ",
    ),
}


class IntegratedFillResult(BaseModel):
    """LLM output: map each template token to its replacement string."""

    replacements: Dict[str, str] = Field(default_factory=dict)


def _deep_truncate(obj: Any, max_str: int = 2200, max_list: int = 60) -> Any:
    if isinstance(obj, dict):
        return {k: _deep_truncate(v, max_str, max_list) for k, v in obj.items()}
    if isinstance(obj, list):
        return [_deep_truncate(i, max_str, max_list) for i in obj[:max_list]]
    if isinstance(obj, str) and len(obj) > max_str:
        return obj[:max_str] + "…"
    return obj


def _overview_to_prompt_json(overview: ProspectOverview) -> str:
    raw = overview.model_dump(mode="json", exclude_none=True)
    trimmed = _deep_truncate(raw)
    try:
        text = json.dumps(trimmed, indent=2, default=str, allow_nan=False)
    except (ValueError, TypeError):
        text = json.dumps(trimmed, indent=2, default=str)
    max_chars = int(os.environ.get("POC_DOC_MAX_PROMPT_CHARS", "75000"))
    if len(text) > max_chars:
        return (
            text[:max_chars]
            + "\n\n…(truncated for model context; prioritize non-null fields above when filling.)\n"
        )
    return text


def _strip_json_fence(text: str) -> str:
    t = text.strip()
    if t.startswith("```"):
        parts = t.split("```")
        if len(parts) >= 2:
            inner = parts[1]
            if inner.lstrip().startswith("json"):
                inner = inner.lstrip()[4:]
            return inner.strip()
    return t


def _extract_json_object(raw: str) -> dict:
    """
    Parse a JSON object from an LLM message that may include prose or markdown fences.
    """
    s = raw.strip().lstrip("\ufeff")
    s = _strip_json_fence(s)
    try:
        data = json.loads(s)
        if isinstance(data, dict):
            return data
    except json.JSONDecodeError:
        pass
    start = s.find("{")
    if start < 0:
        raise ValueError("No JSON object found in model response")
    decoder = json.JSONDecoder()
    try:
        obj, _end = decoder.raw_decode(s[start:])
    except json.JSONDecodeError as e:
        raise ValueError(f"Invalid JSON in model response: {e}") from e
    if not isinstance(obj, dict):
        raise ValueError("Top-level JSON must be an object")
    return obj


def _overview_blob(overview: ProspectOverview) -> str:
    try:
        return json.dumps(overview.model_dump(mode="json", exclude_none=True), default=str).lower()
    except Exception:
        return ""


def _one_line_account_summary(overview: ProspectOverview) -> str:
    bits: list[str] = []
    if overview.salesforce:
        sf = overview.salesforce
        if (sf.name or "").strip():
            bits.append(sf.name.strip())
        if (sf.industry or "").strip():
            bits.append(sf.industry.strip())
    if overview.gong_summary and getattr(overview.gong_summary, "total_calls", None):
        bits.append(f"{overview.gong_summary.total_calls} Gong calls on record")
    if overview.pendo_usage:
        bits.append("Pendo usage on file")
    if overview.product_usage and overview.product_usage.adoption_status:
        bits.append(f"Adoption: {overview.product_usage.adoption_status}")
    base = " · ".join(bits) if bits else (overview.lookup_value or "Account")
    return (base + " — scope and metrics to confirm in discovery.")[:300]


def _first_choice_from_bracket_key(key: str) -> str | None:
    if not (key.startswith("[") and key.endswith("]")):
        return None
    inner = key[1:-1].strip()
    if not inner or inner == "X":
        return None
    parts = [p.strip() for p in inner.split("/") if p.strip()]
    if not parts:
        return None
    return parts[0][:220]


def _keyword_pick_for_bracket(key: str, overview: ProspectOverview) -> str | None:
    """Pick a template option when the key lists vendors/frameworks and CRM text hints at one."""
    blob = _overview_blob(overview)
    if not blob:
        return None
    pairs = [
        ("[OpenAI / Anthropic / AWS Bedrock / GCP Vertex / Azure OAI]", "openai", "OpenAI"),
        ("[OpenAI / Anthropic / AWS Bedrock / GCP Vertex / Azure OAI]", "anthropic", "Anthropic"),
        ("[OpenAI / Anthropic / AWS Bedrock / GCP Vertex / Azure OAI]", "bedrock", "AWS Bedrock"),
        ("[OpenAI / Anthropic / AWS Bedrock / GCP Vertex / Azure OAI]", "vertex", "GCP Vertex"),
        ("[OpenAI / Anthropic / AWS Bedrock / GCP Vertex / Azure OAI]", "azure openai", "Azure OAI"),
        ("[LangChain / LangGraph / LlamaIndex / Custom / etc.]", "langgraph", "LangGraph"),
        ("[LangChain / LangGraph / LlamaIndex / Custom / etc.]", "langchain", "LangChain"),
        ("[LangChain / LangGraph / LlamaIndex / Custom / etc.]", "llamaindex", "LlamaIndex"),
        ("[LangChain / LlamaIndex / LangGraph / Custom / None]", "langgraph", "LangGraph"),
        ("[LangChain / LlamaIndex / LangGraph / Custom / None]", "langchain", "LangChain"),
        ("[LangChain / LlamaIndex / LangGraph / Custom / None]", "llamaindex", "LlamaIndex"),
        ("[DataDog / Splunk / Homegrown / None]", "datadog", "DataDog"),
        ("[DataDog / Splunk / Homegrown / None]", "splunk", "Splunk"),
        ("[Pinecone / Weaviate / pgvector / Internal / N/A]", "pinecone", "Pinecone"),
        ("[Pinecone / Weaviate / pgvector / Internal / N/A]", "weaviate", "Weaviate"),
        ("[Pinecone / Weaviate / pgvector / Internal / N/A]", "pgvector", "pgvector"),
        ("[SaaS / VPC / On-Prem]", "vpc", "VPC"),
        ("[SaaS / VPC / On-Prem]", "on-prem", "On-Prem"),
        ("[SaaS / VPC / On-Prem]", "saas", "SaaS"),
        ("[LLM Calls / RAG / Agentic (Tool Calls) / Multi-Agent]", "multi-agent", "Multi-Agent"),
        ("[LLM Calls / RAG / Agentic (Tool Calls) / Multi-Agent]", "agentic", "Agentic (Tool Calls)"),
        ("[LLM Calls / RAG / Agentic (Tool Calls) / Multi-Agent]", "rag", "RAG"),
        ("[LLM Calls / RAG / Agentic (Tool Calls) / Multi-Agent]", "llm calls", "LLM Calls"),
    ]
    for k_needle, needle, label in pairs:
        if key == k_needle and needle in blob:
            return label
    return None


def _fallback_fill_for_key(
    key: str,
    overview: ProspectOverview,
    document_template: str,
    manual_notes: str | None,
) -> str:
    """Deterministic fill when the LLM omits a key, echoes the placeholder, or returns whitespace."""
    if key == "[Company Name]":
        return _company_display_name(overview)

    if key == "[X]":
        if document_template == "pot":
            return "1"
        if document_template == "poc_saas":
            return "2"
        return "4"

    if key == "___________________________________________":
        return _one_line_account_summary(overview)

    if key == "Start: _____ / End: _____ - [X] Weeks  ":
        # Dates are unknown without SA input — still replace underscores so the line reads as filled.
        if document_template == "pot":
            return "Start: TBD / End: TBD — under 1 week (align in kickoff)  "
        if document_template == "poc_saas":
            return "Start: TBD / End: TBD — target under 2 weeks (align in kickoff)  "
        return "Start: TBD / End: TBD — target under 4 weeks (align in kickoff)  "

    if key.startswith("e.g.,") and ("___" in key or "<" in key):
        return (
            key.replace("___", "TBD")
            .replace("_____", "TBD")
            + " (Refine with customer during scoping.)"
        )[:320]

    picked = _keyword_pick_for_bracket(key, overview)
    if picked:
        return picked

    if key.startswith("[") and key.endswith("]"):
        first = _first_choice_from_bracket_key(key)
        if first:
            return first

    if (manual_notes or "").strip():
        return f"TBD — see SA notes: {(manual_notes or '')[:160]}"

    return "TBD — confirm with account team"


def _coerce_value_for_key(
    key: str,
    raw: Any,
    overview: ProspectOverview,
    document_template: str,
    manual_notes: str | None,
) -> str:
    if key == "[Company Name]":
        return _company_display_name(overview)
    if raw is None:
        return _fallback_fill_for_key(key, overview, document_template, manual_notes)
    s = str(raw).strip()
    if not s or s == key.strip():
        return _fallback_fill_for_key(key, overview, document_template, manual_notes)
    return s[:4000]


def _iter_paragraphs_in_cell(cell: Any) -> Iterable[Paragraph]:
    for p in cell.paragraphs:
        yield p
    for t in cell.tables:
        for row in t.rows:
            for c in row.cells:
                yield from _iter_paragraphs_in_cell(c)


def _iter_body_paragraphs(doc: Document) -> Iterable[Paragraph]:
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                yield from _iter_paragraphs_in_cell(cell)


def _iter_part_paragraphs(part: Any) -> Iterable[Paragraph]:
    if part is None:
        return
    for p in part.paragraphs:
        yield p
    for t in part.tables:
        for row in t.rows:
            for cell in row.cells:
                yield from _iter_paragraphs_in_cell(cell)


def iter_all_paragraphs(doc: Document) -> Iterable[Paragraph]:
    """Body, tables, and header/footer paragraphs."""
    yield from _iter_body_paragraphs(doc)
    for section in doc.sections:
        yield from _iter_part_paragraphs(section.header)
        yield from _iter_part_paragraphs(section.footer)
        if section.different_first_page_header_footer:
            yield from _iter_part_paragraphs(section.first_page_header)
            yield from _iter_part_paragraphs(section.first_page_footer)


def discover_replacement_keys(doc: Document, document_template: str) -> List[str]:
    """Collect bracket tokens plus any configured literal line that exists in this file."""
    found: set[str] = set()
    paras = list(iter_all_paragraphs(doc))
    for p in paras:
        found.update(BRACKET_PLACEHOLDER_RE.findall(p.text))
    for literal in EXTRA_LITERAL_CANDIDATES_BY_TEMPLATE.get(document_template, ()):
        for p in paras:
            if literal in p.text:
                found.add(literal)
                break
    # Longest keys first so shorter tokens (e.g. [X]) do not break longer literals prematurely.
    return sorted(found, key=len, reverse=True)


def _set_paragraph_text_collapsed(paragraph: Paragraph, new_text: str) -> None:
    runs = paragraph.runs
    if runs:
        runs[0].text = new_text
        for r in runs[1:]:
            r.text = ""
    else:
        paragraph.add_run(new_text)


def apply_replacements(doc: Document, mapping: List[Tuple[str, str]]) -> None:
    """Apply (old, new) pairs to every paragraph; longest keys should appear first in the list."""
    for paragraph in iter_all_paragraphs(doc):
        text = paragraph.text
        orig = text
        for old, new in mapping:
            if old in text:
                text = text.replace(old, new)
        if text != orig:
            _set_paragraph_text_collapsed(paragraph, text)


def _set_cell_text(cell: Any, text: str) -> None:
    """Write plain text into the first paragraph of a table cell."""
    if cell.paragraphs:
        _set_paragraph_text_collapsed(cell.paragraphs[0], text)
    else:
        cell.add_paragraph(text)


def _cell_is_effectively_empty(cell: Any) -> bool:
    return not any(p.text.strip() for p in cell.paragraphs)


def _row_data_cells_empty(row: Any) -> bool:
    return all(_cell_is_effectively_empty(c) for c in row.cells)


def _table_header_cells(row: Any) -> List[str]:
    return [c.text.strip() for c in row.cells]


def _first_customer_participant_from_gong(overview: ProspectOverview) -> Tuple[str, str] | None:
    """(name, email) for a likely customer-side Gong participant."""
    gong = overview.gong_summary
    if not gong:
        return None
    for call in gong.recent_calls:
        for p in call.participants:
            name = (p.name or "").strip()
            if not name:
                continue
            aff = (p.affiliation or "").strip().lower()
            if aff in ("company", "internal"):
                continue
            # Gong often uses "non_company" for the prospect; also accept explicit external/customer.
            if aff in ("external", "customer", "prospect", "non_company", ""):
                email = (p.email or "").strip()
                return (name, email)
    return None


def _first_arize_participant_from_gong(overview: ProspectOverview) -> Tuple[str, str, str] | None:
    """(name, email, title_guess) from Gong when affiliation is host company (Arize)."""
    gong = overview.gong_summary
    if not gong:
        return None
    for call in gong.recent_calls:
        for p in call.participants:
            name = (p.name or "").strip()
            if not name:
                continue
            aff = (p.affiliation or "").strip().lower()
            if aff != "company":
                continue
            email = (p.email or "").strip()
            return (name, email, "Arize (from recent Gong)")
    return None


def _arize_roster_fields(sf: Any) -> Tuple[str, str, str] | None:
    """(name, title, role_column) for the primary Arize contact row."""
    sa = (getattr(sf, "assigned_sa", None) or "").strip()
    cse = (getattr(sf, "assigned_cse", None) or "").strip()
    ai = (getattr(sf, "assigned_ai_se", None) or "").strip()
    if sa:
        return (sa, "Solution Architect (Arize)", "Arize")
    if cse:
        return (cse, "Customer Success Engineer (Arize)", "Arize")
    if ai:
        return (ai, "AI Sales Engineer (Arize)", "Arize")
    return None


def _arize_roster_line(overview: ProspectOverview) -> Tuple[str, str, str, str] | None:
    """name, title, role_column, email (may be empty) for the primary Arize contact row."""
    if overview.salesforce:
        line = _arize_roster_fields(overview.salesforce)
        if line:
            nm, title, role = line
            return (nm, title, role, "")
    gong_line = _first_arize_participant_from_gong(overview)
    if gong_line:
        nm, email, title = gong_line
        return (nm, title, "Arize", email)
    if overview.salesforce:
        owner = (getattr(overview.salesforce, "owner_name", None) or "").strip()
        if owner:
            return (owner, "Account Owner (SFDC)", "Arize", "")
    return None


def fill_structured_roster_cells(
    doc: Document, overview: ProspectOverview, document_template: str
) -> None:
    """
    PoC SaaS/VPC templates include roster and kickoff tables with no bracket placeholders.
    Fill the first data row when still blank, using Gong (customer) and Salesforce (Arize).
    """
    if document_template not in ("poc_saas", "poc_vpc") or len(doc.tables) < 4:
        return

    t1 = doc.tables[1]
    if _table_header_cells(t1.rows[0]) == [
        "Name",
        "Role / Team",
        "% Time Committed",
        "Responsibilities",
    ] and _row_data_cells_empty(t1.rows[1]):
        ext = _first_customer_participant_from_gong(overview)
        if ext:
            name, email = ext
            _set_cell_text(t1.rows[1].cells[0], name[:240])
            _set_cell_text(t1.rows[1].cells[1], "Customer team (from recent Gong)")
            _set_cell_text(t1.rows[1].cells[2], "TBD")
            _set_cell_text(t1.rows[1].cells[3], (email or "—")[:240])

    t2 = doc.tables[2]
    if _table_header_cells(t2.rows[0]) == ["Name", "Title", "Email", "Role"] and _row_data_cells_empty(
        t2.rows[1]
    ):
        line = _arize_roster_line(overview)
        if line:
            nm, title, role, email_cell = line
            _set_cell_text(t2.rows[1].cells[0], nm[:240])
            _set_cell_text(t2.rows[1].cells[1], title[:240])
            _set_cell_text(t2.rows[1].cells[2], (email_cell or "")[:240])
            _set_cell_text(t2.rows[1].cells[3], role)

    t3 = doc.tables[3]
    if _table_header_cells(t3.rows[0]) == ["Item", "Details"]:
        arize = _arize_roster_line(overview)
        if arize:
            nm, _, _, _ = arize
            for ri in range(1, len(t3.rows)):
                label = t3.rows[ri].cells[0].text.strip()
                if "Primary Escalation Contact (Arize)" in label and _cell_is_effectively_empty(
                    t3.rows[ri].cells[1]
                ):
                    _set_cell_text(t3.rows[ri].cells[1], nm[:240])
                    break


def _template_instructions(document_template: str) -> str:
    if document_template == "poc_saas":
        return (
            "You are filling an **Arize AX Proof of Concept (PoC) — SaaS** Word template. "
            "Choose concise values that fit table cells. Prefer SaaS deployment language where relevant."
        )
    if document_template == "poc_vpc":
        return (
            "You are filling an **Arize AX Proof of Concept (PoC) — VPC** Word template. "
            "Emphasize customer VPC / private environment, infrastructure alignment, and longer trial windows when appropriate."
        )
    return (
        "You are filling an **Arize AX Proof of Technology (PoT)** Word template. "
        "Emphasize short workshop-style validation, use-case mapping, and discovery outcomes."
    )


def _company_display_name(overview: ProspectOverview) -> str:
    if overview.salesforce and (overview.salesforce.name or "").strip():
        return overview.salesforce.name.strip()
    if (overview.lookup_value or "").strip():
        return overview.lookup_value.strip()
    return "Customer"


def generate_integrated_fills(
    *,
    keys: List[str],
    overview: ProspectOverview,
    document_template: str,
    manual_notes: str | None,
    llm_model: str | None = None,
) -> Dict[str, str]:
    model = llm_model or os.environ.get("POC_DOC_LLM_MODEL", "claude-haiku-4-5")
    payload = _overview_to_prompt_json(overview)
    keys_json = json.dumps(keys, indent=2, ensure_ascii=False)
    notes_block = ""
    if (manual_notes or "").strip():
        notes_block = (
            "\n## Solution Architect notes (use when supported by data; do not invent contradictions)\n"
            f"{manual_notes.strip()}\n"
        )

    prompt = f"""{_template_instructions(document_template)}

You will receive JSON: a **ProspectOverview** from our warehouse (Salesforce, Gong, Pendo, FullStory, etc.).

## Your task
Return **only** a JSON object with a single property `"replacements"` whose keys are **exactly** the strings listed
in KEYS below (same spelling, punctuation, and spacing). Each value replaces that substring wherever it appears in the Word file.

### Rules
- Use only facts grounded in the ProspectOverview JSON (and SA notes when consistent). If unknown, use a short honest placeholder like "TBD — confirm with account team" rather than fabricating metrics.
- **Every key must have a non-empty string value** (never `""`, never null). Do not echo the placeholder key back as the value.
- For keys that look like multiple-choice prompts (e.g. "[OpenAI / Anthropic / ...]"), pick **one** realistic option or a short comma-separated subset that matches the account.
- For bracket keys, keep values **brief** (typically under ~200 characters) so they fit Word table cells.
- For the long `e.g., ...` lines and the `Start: _____ / End: _____ - [X] Weeks  ` line, return a **complete finished line** of similar length (replace the entire template line).
- For `___________________________________________`, replace with **one** short headline-style sentence summarizing the opportunity (no underscores).

## KEYS (required JSON keys under "replacements")
{keys_json}

## ProspectOverview JSON
{payload}
{notes_block}
"""

    response = litellm.completion(
        model=model,
        messages=[{"role": "user", "content": prompt}],
        max_tokens=int(os.environ.get("POC_DOC_MAX_TOKENS", "12288")),
        temperature=0.25,
    )
    raw = (response.choices[0].message.content or "").strip()
    if os.environ.get("POC_DOC_DEBUG", "").strip().lower() in ("1", "true", "yes"):
        print(f"[poc_document_generator] LLM raw (first 800 chars): {raw[:800]!r}")

    try:
        data = _extract_json_object(raw)
        parsed = IntegratedFillResult.model_validate(data)
    except (ValueError, json.JSONDecodeError, ValidationError) as e:
        raise AppendixGenerationError(str(e)) from e

    llm_map = parsed.replacements or {}
    merged: Dict[str, str] = {}
    for k in keys:
        merged[k] = _coerce_value_for_key(
            k,
            llm_map.get(k),
            overview,
            document_template,
            manual_notes,
        )

    return merged


def _safe_filename_part(name: str) -> str:
    s = "".join(c for c in name if c.isalnum() or c in (" ", "-", "_")).strip()
    return s.replace(" ", "_")[:80] or "Account"


def build_poc_document(
    overview: ProspectOverview,
    document_template: str,
    manual_notes: str | None = None,
    llm_model: str | None = None,
) -> tuple[bytes, str]:
    path = TEMPLATES.get(document_template)
    if not path or not path.is_file():
        raise FileNotFoundError(f"Missing template for {document_template}: {path}")

    doc = Document(str(path))
    keys = discover_replacement_keys(doc, document_template)
    if not keys:
        raise ValueError("No replaceable placeholders found in template")

    fill_map = generate_integrated_fills(
        keys=keys,
        overview=overview,
        document_template=document_template,
        manual_notes=manual_notes,
        llm_model=llm_model,
    )
    ordered = sorted(fill_map.items(), key=lambda kv: len(kv[0]), reverse=True)
    apply_replacements(doc, ordered)
    fill_structured_roster_cells(doc, overview, document_template)

    buf = io.BytesIO()
    doc.save(buf)
    date_part = datetime.now(timezone.utc).strftime("%Y%m%d")
    label = overview.lookup_value or ""
    if overview.salesforce and getattr(overview.salesforce, "name", None):
        label = overview.salesforce.name or label
    safe = _safe_filename_part(label or "Account")
    fname = f"Arize_AX_{document_template}_{safe}_{date_part}.docx"
    return buf.getvalue(), fname
